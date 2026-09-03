import importlib.util
import io
import sys
import types
import unittest
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest.mock import patch

from docx import Document as WordDocument


def _load_docx_parser():
    """Load docx_parser without triggering the heavy package __init__.

    ``docreader/parser/__init__.py`` imports doc_parser -> textract, a heavy
    dependency unrelated to DOCX parsing. Registering the packages as bare
    namespaces lets us import only the modules docx_parser actually needs.
    """
    root = Path(__file__).resolve().parents[2]
    docreader_pkg = types.ModuleType("docreader")
    docreader_pkg.__path__ = [str(root / "docreader")]
    sys.modules.setdefault("docreader", docreader_pkg)
    parser_pkg = types.ModuleType("docreader.parser")
    parser_pkg.__path__ = [str(root / "docreader" / "parser")]
    sys.modules["docreader.parser"] = parser_pkg

    spec = importlib.util.spec_from_file_location(
        "docreader.parser.docx_parser", root / "docreader" / "parser" / "docx_parser.py"
    )
    module = importlib.util.module_from_spec(spec)
    sys.modules["docreader.parser.docx_parser"] = module
    spec.loader.exec_module(module)
    return module


docx_parser = _load_docx_parser()
DocxParser = docx_parser.DocxParser
table_to_gfm_markdown = docx_parser.table_to_gfm_markdown


def _parse(content):
    """Parse a DOCX through the real Docx processor.

    The production path uses a ProcessPoolExecutor backed by a multiprocessing
    Manager. Some CI sandboxes forbid POSIX semaphores, so run the same page
    task pool on threads and swap the manager for a plain list - parse
    behavior is identical, only the execution backend changes.
    """

    class _FakeManager:
        def __enter__(self):
            self._items = []
            return self

        def __exit__(self, *exc):
            return False

        def list(self):
            return self._items

    with patch.object(docx_parser, "Manager", _FakeManager), patch.object(
        docx_parser, "ProcessPoolExecutor", ThreadPoolExecutor
    ):
        return DocxParser(max_pages=100).parse_into_text(content)


def _docx_bytes(build):
    doc = WordDocument()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class DocxTableContentTest(unittest.TestCase):
    """Regression test: DOCX tables must be kept in the parsed text.

    Docx.__call__ used to return tables separately from the text lines, and
    parse_into_text dropped them, silently losing all table content.
    """

    def _docx_with_table(self, cells):
        def build(doc):
            doc.add_paragraph("Introduction paragraph")
            table = doc.add_table(rows=len(cells), cols=len(cells[0]))
            for r, row in enumerate(cells):
                for c, value in enumerate(row):
                    table.cell(r, c).text = value

        return _docx_bytes(build)

    def test_table_content_is_kept_in_parsed_text(self):
        content = self._docx_with_table(
            [["City", "Population"], ["Beijing", "21.5M"]]
        )
        document = _parse(content)
        self.assertIn("Introduction paragraph", document.content)
        for cell in ("City", "Population", "Beijing", "21.5M"):
            self.assertIn(cell, document.content)
        self.assertIn("| City | Population |", document.content)
        self.assertIn("| --- | --- |", document.content)

    def test_table_only_document_is_not_empty(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Header"
            table.cell(1, 0).text = "Value"

        document = _parse(_docx_bytes(build))
        self.assertIn("Header", document.content)
        self.assertIn("Value", document.content)

    def test_pipe_in_cell_does_not_break_table(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "A|B"
            table.cell(1, 0).text = "C"

        document = _parse(_docx_bytes(build))
        self.assertIn(r"A\|B", document.content)

    def test_empty_adjacent_cells_keep_columns(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(0, 2).text = "C"
            table.cell(1, 0).text = "1"
            table.cell(1, 1).text = ""
            table.cell(1, 2).text = ""

        document = _parse(_docx_bytes(build))
        self.assertIn("| A | B | C |", document.content)
        self.assertIn("| 1 |  |  |", document.content)

    def test_equal_adjacent_cells_are_not_collapsed(self):
        """Adjacent independent cells with the same text must stay separate.

        Regression for the #2634 control row: ``相同值 | 相同值 | 独立值``.
        """

        def build(doc):
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).text = "相同值"
            table.cell(0, 1).text = "相同值"
            table.cell(0, 2).text = "独立值"
            table.cell(1, 0).text = "x"
            table.cell(1, 1).text = "y"
            table.cell(1, 2).text = "z"

        document = _parse(_docx_bytes(build))
        self.assertIn("| 相同值 | 相同值 | 独立值 |", document.content)
        self.assertIn("| x | y | z |", document.content)

    def test_horizontal_merge_keeps_grid_width(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).merge(table.cell(0, 1)).text = "merged"
            table.cell(0, 2).text = "right"
            table.cell(1, 0).text = "a"
            table.cell(1, 1).text = "b"
            table.cell(1, 2).text = "c"

        document = _parse(_docx_bytes(build))
        self.assertIn("| merged | merged | right |", document.content)
        self.assertIn("| a | b | c |", document.content)

    def test_table_stays_between_paragraphs(self):
        def build(doc):
            doc.add_paragraph("before table")
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "H"
            table.cell(1, 0).text = "V"
            doc.add_paragraph("after table")

        document = _parse(_docx_bytes(build))
        before = document.content.index("before table")
        header = document.content.index("| H |")
        after = document.content.index("after table")
        self.assertLess(before, header)
        self.assertLess(header, after)

    def test_html_like_cell_text_is_literal(self):
        def build(doc):
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "a < b"
            table.cell(0, 1).text = "</td><td>injected"

        document = _parse(_docx_bytes(build))
        self.assertIn("| a < b | </td><td>injected |", document.content)

    def test_table_to_gfm_markdown_helper(self):
        doc = WordDocument()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        self.assertEqual(
            table_to_gfm_markdown(table),
            "| A | B |\n| --- | --- |\n| C | D |",
        )

    def test_table_to_gfm_markdown_skips_empty(self):
        class _Empty:
            rows = []

        self.assertEqual(table_to_gfm_markdown(_Empty()), "")


if __name__ == "__main__":
    unittest.main()
